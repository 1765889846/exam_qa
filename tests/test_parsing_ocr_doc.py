"""PDF OCR 与 .doc 解析测试。"""

from unittest.mock import patch

import pytest

from src.exceptions import BadRequestException
from src.services import parsing
from src.services.parsing import ParsedDocument, ParsedPage, _parse_doc, _parse_pdf


class TestPdfOcr:
    def test_pdf_retries_force_ocr_when_empty(self, monkeypatch):
        monkeypatch.setattr(parsing.config.parsing, "pdf_use_ocr", True)
        monkeypatch.setattr(parsing.config.parsing, "pdf_force_ocr", False)

        empty = ParsedDocument([])
        ocr_doc = ParsedDocument([ParsedPage(page=1, text="OCR 识别文本")])
        calls: list[dict] = []

        def fake_llm(path, *, force_ocr=None):
            calls.append({"force_ocr": force_ocr})
            if force_ocr:
                return ocr_doc
            return empty

        with patch.object(parsing, "_parse_pdf_pymupdf4llm", side_effect=fake_llm):
            with patch.object(parsing, "_parse_pdf_fitz", return_value=empty):
                result = _parse_pdf("scan.pdf")

        assert "OCR 识别文本" in result.full_text
        assert any(c.get("force_ocr") for c in calls)

    def test_pdf_ocr_kwargs_from_config(self, monkeypatch):
        monkeypatch.setattr(parsing.config.parsing, "pdf_use_ocr", True)
        monkeypatch.setattr(parsing.config.parsing, "pdf_force_ocr", True)
        monkeypatch.setattr(parsing.config.parsing, "pdf_ocr_language", "eng+chi_sim")

        kwargs = parsing._pymupdf4llm_kwargs()
        assert kwargs["use_ocr"] is True
        assert kwargs["force_ocr"] is True
        assert kwargs["ocr_language"] == "eng+chi_sim"
        assert kwargs["page_chunks"] is True


class TestParseDoc:
    def test_doc_without_libreoffice_raises(self, temp_dir, monkeypatch):
        fake = temp_dir / "legacy.doc"
        fake.write_bytes(b"fake doc content")
        monkeypatch.setattr(parsing, "_find_soffice", lambda: None)

        with pytest.raises(BadRequestException, match="LibreOffice"):
            _parse_doc(str(fake))

    def test_doc_converts_via_libreoffice(self, temp_dir, monkeypatch):
        from docx import Document

        docx_src = temp_dir / "converted.docx"
        doc = Document()
        doc.add_paragraph("旧版 Word 内容")
        doc.save(str(docx_src))

        fake_doc = temp_dir / "legacy.doc"
        fake_doc.write_bytes(b"placeholder")

        monkeypatch.setattr(parsing, "_find_soffice", lambda: "soffice")
        monkeypatch.setattr(
            parsing,
            "_convert_doc_to_docx",
            lambda path: docx_src if path.endswith("legacy.doc") else None,
        )

        parsed = _parse_doc(str(fake_doc))
        assert "旧版 Word 内容" in parsed.full_text
